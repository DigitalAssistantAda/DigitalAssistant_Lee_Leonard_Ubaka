from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Query
from sqlalchemy.orm import Session
from datetime import datetime, timedelta, timezone
from docx import Document as DocxDocument
from fastapi.responses import PlainTextResponse, StreamingResponse
from io import BytesIO
import hashlib
import os
import re
import uuid
import logging
import requests
from starlette.concurrency import run_in_threadpool
from database import get_db
from models.user import User
from models.document import Document, DocumentStatus
from models.container import Container
from models.workspace import WorkspaceMember, MemberStatus
from models.document_chunk import DocumentChunk
from models.chunk_embedding import ChunkEmbedding
from models.document_deletion_request import DocumentDeletionRequest, DeletionRequestStatus
from schemas.document import DocumentDeletionRequestResponse, DocumentDeletionRequestListResponse
from models.document_hint import DocumentHint
from models.document_duplicate import DocumentDuplicate
from models.summary import Summary
from models.job import Job
from models.embedding_job import EmbeddingJob, EmbeddingJobStatus
from models.audit_log import AuditLog
import math
from sqlalchemy import desc, func
from schemas.document import (
    DocumentResponse,
    DocumentListResponse,
    UpdateDocumentRequest,
    DownloadResponse,
    SmartContainerSuggestionResponse,
    ContainerSuggestionOption,
    AutoOrganizeWorkspaceResponse,
    AutoOrganizedDocumentResult,
    DuplicateUploadCheckResponse,
)
from schemas.auth import SuccessResponse
from utils.auth import get_current_user
from utils.authorization import (
    check_workspace_access,
    check_document_access,
    require_workspace_access,
)
from utils.audit import create_audit_log, AuditActions
from utils.storage import storage
from utils.embeddings import embeddings_service
from config import settings
from tasks.embeddings import process_document_embeddings
from errors import AppError
from realtime import connection_manager
from utils.workspace_members import active_workspace_member_ids
from utils.document_helpers import (
    parse_storage_uri as _parse_storage_uri_impl,
    confidence_label,
    confidence_rank,
    infer_auto_container_name,
    ensure_workspace_container,
    workspace_feedback_boosts,
    delete_document_and_relations,
)

router = APIRouter(tags=["Documents"])

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
# No practical per-container limit; use a high cap to avoid abuse (configurable via env if needed)
MAX_DOCUMENTS_PER_CONTAINER = int(os.getenv("MAX_DOCUMENTS_PER_CONTAINER", "5000"))
ALLOWED_MIME_TYPES = [
    "application/pdf",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
]


def _parse_storage_uri(storage_uri: str) -> tuple[str, str]:
    try:
        return _parse_storage_uri_impl(storage_uri)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Invalid storage URI for document"
        )


async def _notify_documents_changed(
    db: Session,
    workspace_id: int | None,
    actor_user_id: int,
    container_id: int | None,
) -> None:
    if workspace_id is not None:
        user_ids = active_workspace_member_ids(db, workspace_id)
    else:
        user_ids = [actor_user_id]

    payload = {
        "workspace_id": workspace_id,
        "container_id": container_id,
    }
    for user_id in user_ids:
        await connection_manager.send_to_user(
            user_id,
            {
                "type": "documents.changed",
                "payload": payload,
            },
        )
        await connection_manager.send_to_user(
            user_id,
            {
                "type": "containers.changed",
                "payload": {"workspace_id": workspace_id},
            },
        )
        await connection_manager.send_to_user(
            user_id,
            {
                "type": "workspaces.changed",
                "payload": {"workspace_id": workspace_id},
            },
        )


def _sanitize_filename(filename: str) -> str:
    base = os.path.basename(filename or "upload")
    safe = re.sub(r"[^A-Za-z0-9._-]", "_", base)
    return safe or "upload"


def _post_n8n_embedding_trigger(payload: dict) -> None:
    if not settings.n8n_embeddings_trigger_url:
        return
    response = requests.post(
        settings.n8n_embeddings_trigger_url,
        json=payload,
        timeout=10
    )
    if response.status_code >= 400:
        raise RuntimeError(
            f"n8n trigger failed with status code {response.status_code}"
        )


def _validate_upload_content(content: bytes, mime_type: str) -> None:
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty.",
        )

    if mime_type == "application/pdf":
        # PDF files should begin with the %PDF magic header.
        if not content.lstrip().startswith(b"%PDF-"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Uploaded file content is not a valid PDF.",
            )

    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # DOCX files are ZIP containers and should begin with PK signature.
        if not content.startswith(b"PK"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Uploaded file content is not a valid DOCX document.",
            )


def _validate_container_for_workspace(
    db: Session,
    current_user: User,
    workspace_id: int,
    container_id: int,
) -> Container:
    require_workspace_access(workspace_id=workspace_id, user=current_user, db=db)

    container = db.query(Container).filter(
        Container.id == container_id,
        Container.workspace_id == workspace_id,
    ).first()
    if not container:
        raise AppError(
            code="CONTAINER_NOT_ACCESSIBLE",
            message="Container not found or you do not have access.",
            status_code=status.HTTP_404_NOT_FOUND,
        )

    return container


def _validate_container_access(
    db: Session,
    current_user: User,
    container_id: int,
) -> Container:
    container = db.query(Container).filter(Container.id == container_id).first()
    if not container:
        raise AppError(
            code="CONTAINER_NOT_ACCESSIBLE",
            message="Container not found or you do not have access.",
            status_code=status.HTTP_404_NOT_FOUND,
        )

    if container.workspace_id is not None:
        check_workspace_access(current_user, container.workspace_id, db)
    elif container.created_by != current_user.id:
        raise AppError(
            code="CONTAINER_NOT_ACCESSIBLE",
            message="Container not found or you do not have access.",
            status_code=status.HTTP_404_NOT_FOUND,
        )

    return container


def _user_friendly_status(document: Document, job: EmbeddingJob | None) -> tuple[str, str | None]:
    """
    Return (status_label, status_detail) for UI. On-brand, friendly, no internal paths or stack traces.
    """
    status = (document.status.value if hasattr(document.status, "value") else str(document.status or "")).lower()
    err = (job.error_message or "").strip() if job else ""

    if status == "ready":
        return ("Ready to search", None)
    if status == "uploaded":
        if job and getattr(job.status, "value", str(job.status)) == EmbeddingJobStatus.PROCESSING.value:
            if job.total_chunks and job.chunks_processed is not None:
                return ("Indexing…", f"Making it searchable ({job.chunks_processed}/{job.total_chunks})")
            return ("Indexing…", "Making it searchable.")
        return ("In the queue", "Waiting for background worker. Use Restart if it stays here, and ensure the Celery worker is running.")
    if status == "processing":
        if job and job.total_chunks and job.chunks_processed is not None:
            return ("Indexing…", f"Making it searchable ({job.chunks_processed}/{job.total_chunks})")
        return ("Indexing…", "Making it searchable.")
    if status == "failed":
        # Map known backend errors to friendly, safe messages
        err_lower = err.lower()
        fn = (document.filename or "").lower()
        if fn.endswith(".dock"):
            return ("Couldn't index", "The .dock extension isn't supported. If this is a Word doc, save it as .docx and upload again.")
        if "format is not supported" in err_lower or "not supported" in err_lower:
            return ("Couldn't index", "This file type isn't supported yet. Try PDF, DOCX, or plain text.")
        if "no extractable text" in err_lower or "no text" in err_lower:
            return ("Couldn't index", "No text could be read (e.g. image-only or scanned PDF).")
        if "text chunking produced no chunks" in err_lower or "no chunks" in err_lower:
            return ("Couldn't index", "File appears empty or unreadable. Try another file.")
        if "content could not be processed" in err_lower or "could not be processed" in err_lower:
            return ("Couldn't index", "File may be corrupted or in an unexpected format.")
        if "voyage" in err_lower or "api" in err_lower or "timeout" in err_lower or "rate" in err_lower:
            return ("Couldn't index", "Indexing service is busy. Try again in a little bit.")
        if "cannot embed empty" in err_lower or "empty" in err_lower:
            return ("Couldn't index", "No text could be extracted. Try a different file.")
        # Generic fallback — never expose raw error to UI
        return ("Couldn't index", "Something went wrong. Try re-uploading or a different file.")

    return ("Processing", None)


def _suggest_container_name_from_content(document: Document, db: Session) -> str:
    """Suggest a folder name from document content using LLM. Never uses filename-only; returns generic name when content/LLM unavailable."""
    from utils.text_generation import summary_generation_service

    chunks = (
        db.query(DocumentChunk.text)
        .filter(DocumentChunk.document_id == document.id, DocumentChunk.text.isnot(None))
        .order_by(DocumentChunk.chunk_index.asc())
        .limit(6)
        .all()
    )
    combined = " ".join((c[0] or "").strip() for c in chunks).strip()[:4000]
    if not combined:
        return "New folder"
    try:
        if summary_generation_service.is_available():
            return summary_generation_service.suggest_folder_name(combined)
    except Exception:
        logger.warning("LLM folder name suggestion failed for document %s, using generic name", document.id)
    return "New folder"


def _sha256_digest(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


async def _find_exact_duplicate_in_container(
    db: Session,
    container_id: int,
    content_hash: str,
) -> Document | None:
    """Return an existing non-deleted document in the container with identical file bytes."""
    existing_docs = (
        db.query(Document)
        .filter(
            Document.container_id == container_id,
            Document.status != DocumentStatus.DELETED,
        )
        .order_by(Document.created_at.desc())
        .all()
    )

    for existing_doc in existing_docs:
        if not existing_doc.storage_uri:
            continue
        try:
            bucket, path = _parse_storage_uri(existing_doc.storage_uri)
            existing_content = await storage.download(bucket=bucket, path=path)
        except Exception:
            # If storage read fails for one doc, continue checking others.
            continue

        if _sha256_digest(existing_content) == content_hash:
            return existing_doc

    return None


@router.post("/containers/{container_id}/documents/duplicate-check", response_model=DuplicateUploadCheckResponse)
async def check_duplicate_upload_in_container(
    container_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Preflight duplicate detection by exact content match within the target container."""
    _validate_container_access(db, current_user, container_id)

    if not file.content_type or file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not allowed. Allowed types: {', '.join(ALLOWED_MIME_TYPES)}"
        )

    content = await file.read()
    size_bytes = len(content)

    if size_bytes > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )

    _validate_upload_content(content, file.content_type)

    duplicate_doc = await _find_exact_duplicate_in_container(
        db=db,
        container_id=container_id,
        content_hash=_sha256_digest(content),
    )

    if duplicate_doc:
        return DuplicateUploadCheckResponse(
            is_duplicate=True,
            duplicate_document_id=duplicate_doc.id,
            duplicate_filename=duplicate_doc.filename,
            duplicate_created_at=duplicate_doc.created_at,
            message="A matching file already exists in this folder.",
        )

    return DuplicateUploadCheckResponse(
        is_duplicate=False,
        message="No duplicate found in this folder.",
    )


@router.post("/workspaces/{workspace_id}/documents", response_model=DocumentResponse)
async def upload_document(
    workspace_id: int,
    file: UploadFile = File(...),
    container_id: int | None = Form(default=None),
    allow_duplicate: bool = Form(default=False),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Uploads a document to a workspace"""
    
    require_workspace_access(workspace_id=workspace_id, user=current_user, db=db)

    if container_id is not None:
        _validate_container_for_workspace(db, current_user, workspace_id, container_id)
    
    # Validate file
    if not file.content_type or file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not allowed. Allowed types: {', '.join(ALLOWED_MIME_TYPES)}"
        )
    
    # Read file content to check size
    content = await file.read()
    size_bytes = len(content)
    
    if size_bytes > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )

    _validate_upload_content(content, file.content_type)

    if container_id is not None:
        duplicate_doc = await _find_exact_duplicate_in_container(
            db=db,
            container_id=container_id,
            content_hash=_sha256_digest(content),
        )
        if duplicate_doc and not allow_duplicate:
            raise AppError(
                code="DUPLICATE_DOCUMENT_DETECTED",
                message=f'A matching file already exists in this folder ("{duplicate_doc.filename}"). Confirm to upload anyway.',
                status_code=status.HTTP_409_CONFLICT,
            )

    # Check document limit before uploading to avoid wasted storage on rejected requests
    if container_id is not None:
        doc_count = db.query(Document).filter(
            Document.container_id == container_id,
            Document.status != DocumentStatus.DELETED
        ).count()

        if doc_count >= MAX_DOCUMENTS_PER_CONTAINER:
            raise AppError(
                code="MAX_DOCUMENTS_REACHED",
                message=f"You have reached the maximum number of documents ({MAX_DOCUMENTS_PER_CONTAINER}) for this container.",
                status_code=status.HTTP_400_BAD_REQUEST,
            )

    # Upload to storage (S3/MinIO)
    safe_name = _sanitize_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"
    storage_path = f"workspaces/{workspace_id}/documents/{unique_name}"
    try:
        storage_uri = await storage.upload(
            bucket=settings.storage_bucket,
            path=storage_path,
            data=content,
            content_type=file.content_type
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to store document."
        )
    
    # Create document record
    document = Document(
        workspace_id=workspace_id,
        container_id=container_id,
        uploaded_by=current_user.id,
        filename=file.filename,
        mime_type=file.content_type,
        size_bytes=size_bytes,
        storage_uri=storage_uri,
        status=DocumentStatus.UPLOADED
    )
    db.add(document)
    try:
        db.commit()
    except Exception:
        try:
            await storage.delete(bucket=settings.storage_bucket, path=storage_path)
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to save document metadata"
        )
    db.refresh(document)
    
    # Log the upload action
    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_UPLOADED,
        object_type="document",
        object_id=document.id,
        metadata={
            "filename": file.filename,
            "size_bytes": size_bytes,
            "workspace_id": workspace_id
        },
        workspace_id=workspace_id
    )

    # Trigger embedding workflow
    # In development, prefer direct Celery dispatch to avoid external callback dependency.
    if settings.environment.lower() == "development":
        process_document_embeddings.delay(document.id, current_user.id)
    elif settings.n8n_embeddings_trigger_url:
        try:
            await run_in_threadpool(
                _post_n8n_embedding_trigger,
                {
                    "document_id": document.id,
                    "workspace_id": workspace_id,
                    "triggered_by": current_user.id,
                },
            )
        except Exception:
            logger.warning(
                "n8n embedding trigger failed for document_id=%s workspace_id=%s; falling back to celery",
                document.id,
                workspace_id,
            )
            process_document_embeddings.delay(document.id, current_user.id)
    else:
        process_document_embeddings.delay(document.id, current_user.id)
    
    return DocumentResponse.model_validate(document)


@router.post("/containers/{container_id}/documents", response_model=DocumentResponse)
async def upload_document_to_container(
    container_id: int,
    file: UploadFile = File(...),
    allow_duplicate: bool = Form(default=False),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Uploads a document directly to a container (including personal containers)."""

    container = _validate_container_access(db, current_user, container_id)

    # Validate file
    if not file.content_type or file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not allowed. Allowed types: {', '.join(ALLOWED_MIME_TYPES)}"
        )

    # Read file content to check size
    content = await file.read()
    size_bytes = len(content)

    if size_bytes > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )

    _validate_upload_content(content, file.content_type)

    duplicate_doc = await _find_exact_duplicate_in_container(
        db=db,
        container_id=container_id,
        content_hash=_sha256_digest(content),
    )
    if duplicate_doc and not allow_duplicate:
        raise AppError(
            code="DUPLICATE_DOCUMENT_DETECTED",
            message=f'A matching file already exists in this folder ("{duplicate_doc.filename}"). Confirm to upload anyway.',
            status_code=status.HTTP_409_CONFLICT,
        )

    # Enforce per-container document limit
    doc_count = db.query(Document).filter(
        Document.container_id == container_id,
        Document.status != DocumentStatus.DELETED
    ).count()

    if doc_count >= MAX_DOCUMENTS_PER_CONTAINER:
        raise AppError(
            code="MAX_DOCUMENTS_REACHED",
            message=f"You have reached the maximum number of documents ({MAX_DOCUMENTS_PER_CONTAINER}) for this container.",
            status_code=status.HTTP_400_BAD_REQUEST,
        )

    # Upload to storage (S3/MinIO)
    safe_name = _sanitize_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"
    storage_path = f"containers/{container_id}/documents/{unique_name}"
    try:
        storage_uri = await storage.upload(
            bucket=settings.storage_bucket,
            path=storage_path,
            data=content,
            content_type=file.content_type
        )
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to store document."
        )

    document = Document(
        workspace_id=container.workspace_id,
        container_id=container_id,
        uploaded_by=current_user.id,
        filename=file.filename,
        mime_type=file.content_type,
        size_bytes=size_bytes,
        storage_uri=storage_uri,
        status=DocumentStatus.UPLOADED
    )
    db.add(document)

    try:
        db.commit()
    except Exception:
        try:
            await storage.delete(bucket=settings.storage_bucket, path=storage_path)
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to save document metadata"
        )

    db.refresh(document)

    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_UPLOADED,
        object_type="document",
        object_id=document.id,
        metadata={
            "filename": file.filename,
            "size_bytes": size_bytes,
            "workspace_id": container.workspace_id,
            "container_id": container_id,
        },
        workspace_id=container.workspace_id
    )

    if settings.environment.lower() == "development":
        process_document_embeddings.delay(document.id, current_user.id)
    elif settings.n8n_embeddings_trigger_url and container.workspace_id is not None:
        try:
            await run_in_threadpool(
                _post_n8n_embedding_trigger,
                {
                    "document_id": document.id,
                    "workspace_id": container.workspace_id,
                    "triggered_by": current_user.id,
                },
            )
        except Exception:
            logger.warning(
                "n8n embedding trigger failed for document_id=%s container_id=%s; falling back to celery",
                document.id,
                container_id,
            )
            process_document_embeddings.delay(document.id, current_user.id)
    else:
        process_document_embeddings.delay(document.id, current_user.id)

    await _notify_documents_changed(
        db=db,
        workspace_id=document.workspace_id,
        actor_user_id=current_user.id,
        container_id=document.container_id,
    )

    await _notify_documents_changed(
        db=db,
        workspace_id=document.workspace_id,
        actor_user_id=current_user.id,
        container_id=document.container_id,
    )

    return DocumentResponse.model_validate(document)


@router.get("/workspaces/{workspace_id}/documents", response_model=DocumentListResponse)
async def list_documents(
    workspace_id: int,
    container_id: int | None = Query(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lists documents in a workspace"""
    
    check_workspace_access(current_user, workspace_id, db)
    
    # Get documents - filtered by workspace
    query = db.query(Document).filter(
        Document.workspace_id == workspace_id,
        Document.status != DocumentStatus.DELETED
    )

    if container_id is not None:
        _validate_container_for_workspace(db, current_user, workspace_id, container_id)
        query = query.filter(Document.container_id == container_id)

    documents = query.order_by(Document.created_at.desc()).all()
    doc_ids = [d.id for d in documents]
    job_by_doc = {}
    if doc_ids:
        jobs = (
            db.query(EmbeddingJob)
            .filter(EmbeddingJob.document_id.in_(doc_ids))
            .order_by(desc(EmbeddingJob.created_at))
            .all()
        )
        for j in jobs:
            if j.document_id not in job_by_doc:
                job_by_doc[j.document_id] = j
    items = []
    for d in documents:
        r = DocumentResponse.model_validate(d)
        r.status_label, r.status_detail = _user_friendly_status(d, job_by_doc.get(d.id))
        items.append(r)
    return DocumentListResponse(items=items)


@router.get("/containers/{container_id}/documents", response_model=DocumentListResponse)
async def list_documents_for_container(
    container_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    container = _validate_container_access(db, current_user, container_id)

    query = db.query(Document).filter(
        Document.container_id == container_id,
        Document.status != DocumentStatus.DELETED,
    )

    if container.workspace_id is not None:
        query = query.filter(Document.workspace_id == container.workspace_id)

    documents = query.order_by(Document.created_at.desc()).all()
    doc_ids = [d.id for d in documents]
    job_by_doc = {}
    if doc_ids:
        jobs = (
            db.query(EmbeddingJob)
            .filter(EmbeddingJob.document_id.in_(doc_ids))
            .order_by(desc(EmbeddingJob.created_at))
            .all()
        )
        for j in jobs:
            if j.document_id not in job_by_doc:
                job_by_doc[j.document_id] = j
    items = []
    for d in documents:
        r = DocumentResponse.model_validate(d)
        r.status_label, r.status_detail = _user_friendly_status(d, job_by_doc.get(d.id))
        items.append(r)
    return DocumentListResponse(items=items)


@router.get("/documents/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Retrieves document metadata"""
    
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    job = (
        db.query(EmbeddingJob)
        .filter(EmbeddingJob.document_id == document_id)
        .order_by(desc(EmbeddingJob.created_at))
        .first()
    )
    r = DocumentResponse.model_validate(document)
    r.status_label, r.status_detail = _user_friendly_status(document, job)
    return r


@router.get("/documents/{document_id}/suggest-container", response_model=SmartContainerSuggestionResponse)
async def suggest_document_container(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Suggest the best workspace container for a document using chunk similarity with safe fallbacks."""
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    if document.workspace_id is None:
        return SmartContainerSuggestionResponse(
            document_id=document.id,
            confidence="low",
            confidence_score=None,
            boost_applied=False,
            reason="Smart organization is available for workspace documents only.",
            alternatives=[],
            suggested_new_container_name=_suggest_container_name_from_content(document, db),
        )

    check_workspace_access(current_user, document.workspace_id, db)

    workspace_containers = db.query(Container).filter(
        Container.workspace_id == document.workspace_id,
    ).all()

    if not workspace_containers:
        return SmartContainerSuggestionResponse(
            document_id=document.id,
            confidence="low",
            confidence_score=None,
            boost_applied=False,
            reason="No destination folders exist in this workspace yet.",
            alternatives=[],
            suggested_new_container_name=_suggest_container_name_from_content(document, db),
        )

    if document.status != DocumentStatus.READY:
        return SmartContainerSuggestionResponse(
            document_id=document.id,
            confidence="low",
            confidence_score=None,
            boost_applied=False,
            reason="Document is still processing. Try again when indexing is complete.",
            alternatives=[],
            suggested_new_container_name=_suggest_container_name_from_content(document, db),
        )

    container_by_id = {container.id: container for container in workspace_containers}
    alternatives: list[ContainerSuggestionOption] = []
    best_container_id: int | None = None
    best_score = 0.0
    fallback_used = False
    feedback_boosts = workspace_feedback_boosts(db, document.workspace_id)

    try:
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document.id,
        ).order_by(DocumentChunk.chunk_index.asc()).limit(4).all()

        combined_text = "\n".join((chunk.text or "") for chunk in chunks).strip()
        if combined_text:
            query_embedding = embeddings_service.generate_embedding(combined_text[:6000])
            similar_rows = embeddings_service.find_similar_embeddings(
                query_embedding=query_embedding,
                workspace_id=document.workspace_id,
                limit=120,
                threshold=0.2,
                db=db,
            )

            max_similarity_by_container: dict[int, float] = {}
            for _chunk_id, similar_doc_id, similarity in similar_rows:
                if similar_doc_id == document.id:
                    continue
                similar_doc = db.query(Document).filter(
                    Document.id == similar_doc_id,
                    Document.workspace_id == document.workspace_id,
                    Document.status != DocumentStatus.DELETED,
                ).first()
                if not similar_doc or not similar_doc.container_id:
                    continue
                if similar_doc.container_id not in container_by_id:
                    continue
                sim = float(similarity)
                max_similarity_by_container[similar_doc.container_id] = max(
                    max_similarity_by_container.get(similar_doc.container_id, 0.0),
                    sim,
                )

            adjusted_scores = {
                container_id: score + feedback_boosts.get(container_id, 0.0)
                for container_id, score in max_similarity_by_container.items()
            }

            ranked = sorted(adjusted_scores.items(), key=lambda item: item[1], reverse=True)
            top = ranked[:3]
            if top:
                best_container_id, best_score = top[0]
                alternatives = [
                    ContainerSuggestionOption(
                        container_id=container_id,
                        container_name=container_by_id[container_id].name,
                        score=round(score, 3),
                    )
                    for container_id, score in top
                ]
    except Exception as exc:
        logger.warning("Container suggestion similarity failed for document %s: %s", document.id, exc)

    # Only suggest an existing container when we have embedding-based semantic similarity (actual ML).
    # We do not use keyword or count-based fallbacks.
    if not best_container_id:
        return SmartContainerSuggestionResponse(
            document_id=document.id,
            confidence="low",
            confidence_score=None,
            boost_applied=False,
            reason="No strong destination signal found. Create a new subfolder or pick one manually.",
            alternatives=[],
            suggested_new_container_name=_suggest_container_name_from_content(document, db),
        )

    best_container = container_by_id[best_container_id]
    boost_applied = feedback_boosts.get(best_container_id, 0.0) > 0
    reason = (
        f"Suggested based on semantic similarity to other files in '{best_container.name}'."
        if not fallback_used
        else f"Suggested by workspace usage pattern: '{best_container.name}' has the strongest current grouping."
    )

    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_CONTAINER_SUGGESTED,
        object_type="document",
        object_id=document.id,
        metadata={
            "workspace_id": document.workspace_id,
            "suggested_container_id": best_container_id,
            "suggested_container_name": best_container.name,
            "confidence": confidence_label(best_score),
            "fallback_used": fallback_used,
        },
        workspace_id=document.workspace_id,
    )

    suggested_new = _suggest_container_name_from_content(document, db)
    return SmartContainerSuggestionResponse(
        document_id=document.id,
        suggested_container_id=best_container_id,
        suggested_container_name=best_container.name,
        suggested_new_container_name=suggested_new,
        confidence=confidence_label(best_score),
        confidence_score=round(float(best_score), 3),
        boost_applied=boost_applied,
        reason=reason,
        alternatives=alternatives,
    )


@router.post("/workspaces/{workspace_id}/documents/auto-organize", response_model=AutoOrganizeWorkspaceResponse)
async def auto_organize_workspace_documents(
    workspace_id: int,
    min_confidence: str = Query(default="high"),
    dry_run: bool = Query(default=False),
    auto_create_missing: bool = Query(default=True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Autonomously move ready workspace documents into suggested containers."""
    require_workspace_access(workspace_id=workspace_id, user=current_user, db=db)

    target_rank = confidence_rank(min_confidence)
    docs = db.query(Document).filter(
        Document.workspace_id == workspace_id,
        Document.status == DocumentStatus.READY,
        Document.status != DocumentStatus.DELETED,
    ).order_by(Document.created_at.desc()).all()

    considered = 0
    moved = 0
    skipped_low_confidence = 0
    skipped_no_suggestion = 0
    skipped_already_organized = 0
    moved_documents: list[AutoOrganizedDocumentResult] = []

    for doc in docs:
        considered += 1
        suggestion = await suggest_document_container(doc.id, current_user, db)
        suggested_container_id = suggestion.suggested_container_id

        if not suggested_container_id:
            if not auto_create_missing:
                skipped_no_suggestion += 1
                continue

            inferred_name = (suggestion.suggested_new_container_name or "").strip() or _suggest_container_name_from_content(doc, db)
            target_container_id = None
            target_container_name = inferred_name
            if not dry_run:
                inferred_container, was_created = ensure_workspace_container(
                    db=db,
                    workspace_id=workspace_id,
                    name=inferred_name,
                    actor_user_id=current_user.id,
                )
                if was_created:
                    create_audit_log(
                        db,
                        current_user,
                        action="container.created_by_auto_organize",
                        object_type="container",
                        object_id=inferred_container.id,
                        metadata={
                            "workspace_id": workspace_id,
                            "name": inferred_container.name,
                            "trigger": "auto_organize",
                        },
                        workspace_id=workspace_id,
                    )
                target_container_id = inferred_container.id
                target_container_name = inferred_container.name

            moved_documents.append(
                AutoOrganizedDocumentResult(
                    document_id=doc.id,
                    filename=doc.filename,
                    from_container_id=doc.container_id,
                    to_container_id=target_container_id,
                    to_container_name=target_container_name,
                    confidence="low",
                    confidence_score=None,
                    boost_applied=False,
                )
            )

            if not dry_run and target_container_id is not None:
                await update_document(
                    document_id=doc.id,
                    request=UpdateDocumentRequest(container_id=target_container_id),
                    current_user=current_user,
                    db=db,
                )
                moved += 1
            continue

        if confidence_rank(suggestion.confidence) < target_rank:
            skipped_low_confidence += 1
            continue

        if doc.container_id == suggested_container_id:
            skipped_already_organized += 1
            continue

        moved_documents.append(
            AutoOrganizedDocumentResult(
                document_id=doc.id,
                filename=doc.filename,
                from_container_id=doc.container_id,
                to_container_id=suggested_container_id,
                to_container_name=suggestion.suggested_container_name,
                confidence=suggestion.confidence,
                confidence_score=suggestion.confidence_score,
                boost_applied=suggestion.boost_applied,
            )
        )

        if not dry_run:
            await update_document(
                document_id=doc.id,
                request=UpdateDocumentRequest(container_id=suggested_container_id),
                current_user=current_user,
                db=db,
            )
            moved += 1

    return AutoOrganizeWorkspaceResponse(
        workspace_id=workspace_id,
        considered=considered,
        moved=moved if not dry_run else len(moved_documents),
        skipped_low_confidence=skipped_low_confidence,
        skipped_no_suggestion=skipped_no_suggestion,
        skipped_already_organized=skipped_already_organized,
        dry_run=dry_run,
        moved_documents=moved_documents,
    )


@router.put("/documents/{document_id}", response_model=DocumentResponse)
async def update_document(
    document_id: int,
    request: UpdateDocumentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Updates document metadata"""

    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    old_filename = document.filename
    old_container_id = document.container_id

    if request.filename:
        document.filename = request.filename

    if request.container_id is not None:
        _validate_container_for_workspace(db, current_user, document.workspace_id, request.container_id)
        document.container_id = request.container_id

    db.commit()
    db.refresh(document)

    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_UPDATED,
        object_type="document",
        object_id=document.id,
        metadata={
            "old_filename": old_filename,
            "new_filename": document.filename,
            "old_container_id": old_container_id,
            "new_container_id": document.container_id,
            "workspace_id": document.workspace_id,
        },
        workspace_id=document.workspace_id,
    )

    if request.container_id is not None and old_container_id != document.container_id:
        metadata = {
            "workspace_id": document.workspace_id,
            "old_container_id": old_container_id,
            "new_container_id": document.container_id,
        }
        if request.suggested_container_id is not None and request.suggested_container_id != document.container_id:
            metadata["suggested_container_id"] = request.suggested_container_id
            metadata["corrected"] = True
        create_audit_log(
            db,
            current_user,
            action=AuditActions.DOCUMENT_CONTAINER_SUGGESTION_APPLIED,
            object_type="document",
            object_id=document.id,
            metadata=metadata,
            workspace_id=document.workspace_id,
        )

    await _notify_documents_changed(
        db=db,
        workspace_id=document.workspace_id,
        actor_user_id=current_user.id,
        container_id=document.container_id,
    )

    return DocumentResponse.model_validate(document)


@router.post("/documents/{document_id}/retry-indexing", response_model=DocumentResponse)
async def retry_document_indexing(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Re-run indexing for a failed or stuck document. Clears existing chunks/embeddings and re-queues the job."""
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    chunk_ids = [row[0] for row in db.query(DocumentChunk.id).filter(
        DocumentChunk.document_id == document_id
    ).all()]
    if chunk_ids:
        db.query(ChunkEmbedding).filter(
            ChunkEmbedding.chunk_id.in_(chunk_ids)
        ).delete(synchronize_session=False)
    db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete(synchronize_session=False)
    db.query(DocumentHint).filter(DocumentHint.document_id == document_id).delete(synchronize_session=False)
    db.query(Job).filter(Job.document_id == document_id).delete(synchronize_session=False)
    db.query(EmbeddingJob).filter(EmbeddingJob.document_id == document_id).delete(synchronize_session=False)
    db.query(Summary).filter(
        Summary.document_id == document_id
    ).update({Summary.document_id: None}, synchronize_session=False)
    db.query(DocumentDuplicate).filter(
        DocumentDuplicate.duplicate_of_id == document_id
    ).update({DocumentDuplicate.duplicate_of_id: None}, synchronize_session=False)
    db.query(DocumentDuplicate).filter(
        DocumentDuplicate.document_id == document_id
    ).delete(synchronize_session=False)

    document.status = DocumentStatus.UPLOADED
    db.commit()
    db.refresh(document)

    create_audit_log(
        db,
        current_user,
        action="document.retry_indexing_requested",
        object_type="document",
        object_id=document.id,
        metadata={
            "filename": document.filename,
            "workspace_id": document.workspace_id,
            "container_id": document.container_id,
        },
        workspace_id=document.workspace_id,
    )

    process_document_embeddings.delay(document_id, current_user.id)

    await _notify_documents_changed(
        db=db,
        workspace_id=document.workspace_id,
        actor_user_id=current_user.id,
        container_id=document.container_id,
    )

    r = DocumentResponse.model_validate(document)
    r.status_label, r.status_detail = _user_friendly_status(document, None)
    return r


@router.get("/documents/{document_id}/download", response_model=DownloadResponse)
async def download_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate a time-limited signed URL for downloading a document"""

    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    if not document.storage_uri:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Document does not have a valid storage location. It may still be initializing.",
        )

    ttl_seconds = max(60, min(int(settings.download_url_ttl_seconds), 7 * 24 * 60 * 60))

    try:
        bucket, path = _parse_storage_uri(document.storage_uri)
        signed_url = await storage.create_download_url(
            bucket=bucket,
            path=path,
            expires_seconds=ttl_seconds,
            filename=document.filename,
            content_type=document.mime_type,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to generate signed download URL for document %s: %s", document.id, e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate download URL"
        )

    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_DOWNLOADED,
        object_type="document",
        object_id=document.id,
        metadata={
            "filename": document.filename,
            "size_bytes": document.size_bytes,
            "workspace_id": document.workspace_id
        },
        workspace_id=document.workspace_id
    )

    expires_at = datetime.now(timezone.utc) + timedelta(seconds=ttl_seconds)
    return DownloadResponse(url=signed_url, expires_at=expires_at)

@router.get("/documents/{document_id}/content")
async def get_document_content(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get raw text content from TXT or DOCX files"""
    
    doc = check_document_access(current_user, document_id, db)
    if doc.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        bucket, path = _parse_storage_uri(doc.storage_uri)
        file_content = await storage.download(bucket=bucket, path=path)
        
        if doc.filename.lower().endswith('.txt'):
            content = file_content.decode('utf-8')
            return PlainTextResponse(content)
        
        elif doc.filename.lower().endswith('.docx'):
            from io import BytesIO
            # Convert bytes to file-like object
            file_stream = BytesIO(file_content)
            docx_doc = DocxDocument(file_stream)
            content = '\n'.join([paragraph.text for paragraph in docx_doc.paragraphs])
            return PlainTextResponse(content)
        
        else:
            raise HTTPException(status_code=400, detail="File type not supported for preview")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reading document {document_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Error reading file")

@router.get("/documents/{document_id}/preview")
async def preview_document(
    document_id: int,
    token: str = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Preview document - returns file for iframe"""
    
    doc = check_document_access(current_user, document_id, db)
    if doc.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        bucket, path = _parse_storage_uri(doc.storage_uri)
        file_content = await storage.download(bucket=bucket, path=path)
        
        return StreamingResponse(
            iter([file_content]),
            media_type=doc.mime_type,
            headers={
                "Content-Disposition": f"inline; filename={doc.filename}",
                "Cache-Control": "no-cache",
                "Content-Type": doc.mime_type,
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error previewing document {document_id}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="Error reading file")
    
# NEW ENDPOINT: Request deletion (if not owner)
@router.post("/documents/{document_id}/deletion-request", response_model=DocumentDeletionRequestResponse)
async def request_document_deletion(
    document_id: int,
    reason: str | None = Query(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Request approval to delete another user's document"""
    
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Can't request deletion of own document
    if document.uploaded_by == current_user.id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="You can delete your own documents directly"
        )
    
    # Check if request already pending
    existing = db.query(DocumentDeletionRequest).filter(
        DocumentDeletionRequest.document_id == document_id,
        DocumentDeletionRequest.requested_by == current_user.id,
        DocumentDeletionRequest.status == DeletionRequestStatus.PENDING
    ).first()
    
    if existing:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="You already have a pending deletion request for this document"
        )
    
    # Create deletion request
    deletion_request = DocumentDeletionRequest(
        document_id=document_id,
        requested_by=current_user.id,
        document_owner=document.uploaded_by,
        reason=reason
    )
    db.add(deletion_request)
    db.commit()
    db.refresh(deletion_request)
    
    # Log action
    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_DELETION_REQUESTED,
        object_type="document",
        object_id=document_id,
        metadata={
            "document_name": document.filename,
            "reason": reason or "No reason provided",
            "workspace_id": document.workspace_id
        },
        workspace_id=document.workspace_id
    )
    
    return DocumentDeletionRequestResponse.model_validate(deletion_request)


# NEW ENDPOINT: Get deletion requests for a document
@router.get("/documents/{document_id}/deletion-requests", response_model=DocumentDeletionRequestListResponse)
async def get_deletion_requests(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get pending deletion requests for a document (owner only)"""
    
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Only document owner can see requests
    if document.uploaded_by != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the document owner can view deletion requests"
        )
    
    requests = db.query(DocumentDeletionRequest).filter(
        DocumentDeletionRequest.document_id == document_id
    ).order_by(DocumentDeletionRequest.created_at.desc()).all()
    
    return DocumentDeletionRequestListResponse(
        items=[DocumentDeletionRequestResponse.model_validate(r) for r in requests]
    )


# NEW ENDPOINT: Approve deletion request
@router.post("/documents/{document_id}/deletion-requests/{request_id}/approve", response_model=DocumentDeletionRequestResponse)
async def approve_deletion_request(
    document_id: int,
    request_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Approve a deletion request and delete the document"""
    
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Only owner can approve
    if document.uploaded_by != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the document owner can approve deletion requests"
        )
    
    deletion_request = db.query(DocumentDeletionRequest).filter(
        DocumentDeletionRequest.id == request_id,
        DocumentDeletionRequest.document_id == document_id
    ).first()
    
    if not deletion_request:
        raise HTTPException(status_code=404, detail="Deletion request not found")
    
    if deletion_request.status != DeletionRequestStatus.PENDING:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Request is already {deletion_request.status}"
        )
    
    # Update request status — committed atomically with document deletion below
    deletion_request.status = DeletionRequestStatus.APPROVED
    deletion_request.responded_at = datetime.now(timezone.utc)

    requester = db.query(User).filter(User.id == deletion_request.requested_by).first()
    try:
        workspace_id, container_id, filename = await delete_document_and_relations(db, document)
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Invalid document storage"
        )
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )

    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_DELETION_APPROVED,
        object_type="document",
        object_id=document_id,
        metadata={
            "document_name": filename,
            "approved_for_user": requester.username if requester else "Unknown",
            "workspace_id": workspace_id
        },
        workspace_id=workspace_id
    )

    await _notify_documents_changed(
        db=db,
        workspace_id=workspace_id,
        actor_user_id=current_user.id,
        container_id=container_id,
    )

    return DocumentDeletionRequestResponse.model_validate(deletion_request)


# NEW ENDPOINT: Deny deletion request
@router.post("/documents/{document_id}/deletion-requests/{request_id}/deny", response_model=DocumentDeletionRequestResponse)
async def deny_deletion_request(
    document_id: int,
    request_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Deny a deletion request"""
    
    document = check_document_access(current_user, document_id, db)
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Only owner can deny
    if document.uploaded_by != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the document owner can deny deletion requests"
        )
    
    deletion_request = db.query(DocumentDeletionRequest).filter(
        DocumentDeletionRequest.id == request_id,
        DocumentDeletionRequest.document_id == document_id
    ).first()
    
    if not deletion_request:
        raise HTTPException(status_code=404, detail="Deletion request not found")
    
    if deletion_request.status != DeletionRequestStatus.PENDING:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Request is already {deletion_request.status}"
        )
    
    deletion_request.status = DeletionRequestStatus.DENIED
    deletion_request.responded_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(deletion_request)
    
    # Log action
    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_DELETION_DENIED,
        object_type="document",
        object_id=document_id,
        metadata={
            "document_name": document.filename,
            "denied_for_user_id": deletion_request.requested_by,
            "workspace_id": document.workspace_id
        },
        workspace_id=document.workspace_id
    )
    
    return DocumentDeletionRequestResponse.model_validate(deletion_request)    

@router.delete("/documents/{document_id}", response_model=SuccessResponse)
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Permanently deletes a document (owner only)"""
    
    document = check_document_access(current_user, document_id, db)

    # Return 404 before 403 to avoid revealing that a deleted document once existed
    if document.status == DocumentStatus.DELETED:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Only owner can delete directly
    if document.uploaded_by != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the document owner can delete. Request approval from the owner."
        )

    try:
        workspace_id, container_id, filename = await delete_document_and_relations(db, document)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Invalid document storage"
        )
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to permanently delete document: {str(e)}"
        )

    create_audit_log(
        db,
        current_user,
        action=AuditActions.DOCUMENT_DELETED,
        object_type="document",
        object_id=document_id,
        metadata={
            "filename": filename,
            "workspace_id": workspace_id
        },
        workspace_id=workspace_id
    )

    await _notify_documents_changed(
        db=db,
        workspace_id=workspace_id,
        actor_user_id=current_user.id,
        container_id=container_id,
    )
    
    return SuccessResponse()


